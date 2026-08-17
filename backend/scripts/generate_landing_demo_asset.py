"""One-off generator for the landing page's real demo example — NOT a recurring
operational command, so it isn't wired into the Makefile.

Runs the real pipeline nodes directly (transcribe -> extract -> validate -> render ->
convert to PDF) against a real audio file, bypassing the DB entirely: node functions are
called via `.__wrapped__` (the raw function `observed_node` preserves via functools.wraps)
so no tenant/report/execution_log rows are needed. Reuses the same field schema and
auto-generated .docx template as the seeded demo tenant.

Usage (run inside the backend container so GROQ/ANTHROPIC keys and GOTENBERG_URL resolve):
    docker compose --project-directory . -f infra/docker-compose.yml exec backend \
        python scripts/generate_landing_demo_asset.py scripts/_landing_demo_input.mp3

Writes to backend/scripts/_landing_demo_output/ (or _landing_demo_output_en/ when the
input filename ends in "_en"): audio.mp3, transcript.txt, fields.json, informe.pdf,
meta.json (cost/latency) — copy the ones you want into frontend/public/demo/.
"""

import asyncio
import json
import shutil
import sys
import uuid
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.core.config import settings
from app.services.agent.nodes.extract import extract_node, validate_node
from app.services.agent.nodes.media import transcribe_node
from app.services.agent.state import AgentState
from app.services.rendering.docx_render import fill_template
from app.services.rendering.gotenberg_client import convert_docx_to_pdf
from scripts.seed_demo_tenant import (
    FIELD_SCHEMA,
    PROMPT_INSTRUCTIONS,
    _GRAY,
    _NAVY,
    _generate_template,
    _heading,
    _set_cell_shading,
    _tag_paragraph,
)


def _base_state(audio_path: Path) -> AgentState:
    return AgentState(
        thread_id="landing-demo",
        tenant_id=uuid.uuid4(),
        channel_connection_id=uuid.uuid4(),
        channel_type="landing_demo",
        sender_id="landing-demo",
        report_id=uuid.uuid4(),
        raw_payload={},
        media_local_path=str(audio_path),
        document_type_id=uuid.uuid4(),
        document_type_name="Meeting Minutes",
        field_schema=FIELD_SCHEMA,
        prompt_instructions=PROMPT_INSTRUCTIONS,
    )


def _generate_template_en(output_path: Path) -> None:
    """English counterpart to seed_demo_tenant._generate_template — kept separate
    because that one is also used to seed the product's real (Spanish) demo tenant;
    this landing-only variant must not change its output."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title_p = document.add_paragraph()
    title_run = title_p.add_run("MEETING MINUTES")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = _NAVY
    subtitle_p = document.add_paragraph()
    subtitle_run = subtitle_p.add_run("{{ company_name }}")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = _GRAY
    subtitle_p.paragraph_format.space_after = Pt(12)

    meta_p = document.add_paragraph()
    meta_p.add_run("Date: ").bold = True
    meta_p.add_run("{{ meeting_date }}    ")
    meta_p.add_run("Location: ").bold = True
    meta_p.add_run("{{ location }}")

    _heading(document, "Attendees", 13)
    _tag_paragraph(document, "{%p for a in attendees %}")
    _tag_paragraph(document, "{{ a }}", style="List Bullet")
    _tag_paragraph(document, "{%p endfor %}")

    _heading(document, "Agenda", 13)
    _tag_paragraph(document, "{%p for item in agenda_items %}")
    _tag_paragraph(document, "{{ item }}", style="List Number")
    _tag_paragraph(document, "{%p endfor %}")

    _heading(document, "Summary", 13)
    document.add_paragraph("{{ summary }}")

    _heading(document, "Decisions Made", 13)
    _tag_paragraph(document, "{%p for d in decisions %}")
    _tag_paragraph(document, "{{ d }}", style="List Bullet")
    _tag_paragraph(document, "{%p endfor %}")

    _heading(document, "Action Items & Owners", 13)
    table = document.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, ["Action", "Owner", "Due Date"], strict=True):
        _set_cell_shading(cell, "1F2A44")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    table.rows[1].cells[0].paragraphs[0].add_run("{%tr for item in action_items %}")
    content_cells = table.rows[2].cells
    content_cells[0].paragraphs[0].add_run("{{ item }}")
    content_cells[1].paragraphs[0].add_run("{{ action_owners[loop.index0] }}")
    content_cells[2].paragraphs[0].add_run("{{ action_due_dates[loop.index0] }}")
    table.rows[3].cells[0].paragraphs[0].add_run("{%tr endfor %}")

    _tag_paragraph(document, "{%p if next_meeting %}")
    next_meeting_p = document.add_paragraph()
    next_meeting_p.add_run("Next Meeting: ").bold = True
    next_meeting_p.add_run("{{ next_meeting }}")
    _tag_paragraph(document, "{%p endif %}")

    footer_p = document.sections[0].footer.paragraphs[0]
    footer_run = footer_p.add_run("Automatically generated by ReportAI — confidential")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = _GRAY

    document.save(str(output_path))


async def generate(audio_path: Path) -> None:
    locale = "en" if audio_path.stem.endswith("_en") else "es"
    output_dir = Path(__file__).parent / (f"_landing_demo_output_{locale}" if locale == "en" else "_landing_demo_output")
    template_fn = _generate_template_en if locale == "en" else _generate_template
    output_dir.mkdir(exist_ok=True)
    state = _base_state(audio_path)

    # settings.TRANSCRIPTION_LANGUAGE defaults to "es" (the product's real market) —
    # this process-local override only affects this one-off script run, not the
    # running backend server, and keeps Whisper from drifting into Spanish mid-way
    # through an English recording.
    settings.TRANSCRIPTION_LANGUAGE = locale

    print("Transcribing (Groq Whisper, real call)...")
    # observed_node wraps with functools.wraps, which sets __wrapped__ to the raw node —
    # bypassing DB-backed execution logging, which this standalone script has no rows for.
    state = await transcribe_node.__wrapped__(state)  # type: ignore[attr-defined]
    assert state.transcript
    print(f"  -> {state.transcript}")

    print("Extracting fields (Claude, real call)...")
    state = await extract_node.__wrapped__(state)  # type: ignore[attr-defined]
    print(f"  -> {state.extracted_fields}")

    print("Validating against the field schema...")
    state = await validate_node.__wrapped__(state)  # type: ignore[attr-defined]
    if state.last_validation_error:
        sys.exit(f"Validation failed, fix the transcript or schema:\n{state.last_validation_error}")

    print("Rendering .docx from the auto-generated template...")
    template_path = output_dir / "template.docx"
    template_fn(template_path)
    docx_path = output_dir / "informe.docx"
    await asyncio.to_thread(fill_template, str(template_path), state.extracted_fields, str(docx_path))

    print("Converting to PDF via Gotenberg...")
    pdf_path = output_dir / "informe.pdf"
    await convert_docx_to_pdf(str(docx_path), str(pdf_path))

    shutil.copy(audio_path, output_dir / "audio.mp3")
    (output_dir / "transcript.txt").write_text(state.transcript, encoding="utf-8")
    (output_dir / "fields.json").write_text(
        json.dumps(state.extracted_fields, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    usage = state.last_tool_usage
    meta = {
        "generated_at": datetime.now(UTC).isoformat(),
        "extraction_model": usage.model_used if usage else None,
        "extraction_cost_usd": usage.cost_usd if usage else None,
    }
    (output_dir / "meta.json").write_text(json.dumps(meta, indent=2), encoding="utf-8")

    print(f"\nDone. Artifacts in {output_dir}")
    print(f"Extraction cost: ${meta['extraction_cost_usd']:.6f}" if meta["extraction_cost_usd"] else "")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        sys.exit("Usage: python scripts/generate_landing_demo_asset.py <path-to-audio-file>")
    asyncio.run(generate(Path(sys.argv[1])))
