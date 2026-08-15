"""Seeds a demo tenant with a self-generated document template — no manual asset creation.
Idempotent: safe to run repeatedly, does nothing if the demo tenant already exists.

Usage: make seed-demo            # seed if missing
       make reset-demo           # DELETE the demo tenant (cascades) and re-seed clean

Credentials come from DEMO_USER_EMAIL / DEMO_USER_PASSWORD / DEMO_TELEGRAM_BOT_TOKEN
when set (prod), with dev-only fallbacks otherwise.
"""

import asyncio
import secrets
import sys
import uuid
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.database import AsyncSessionLocal
from app.core.security import hash_password
from app.models.channel_connection import ChannelConnection
from app.models.document_template import DocumentTemplate
from app.models.document_type import DocumentType
from app.models.tenant import Tenant
from app.models.tenant_user import TenantUser
from app.repositories.base import BaseRepository

DEMO_USER_EMAIL = settings.DEMO_USER_EMAIL or "demo@reportai.dev"
DEMO_USER_PASSWORD = settings.DEMO_USER_PASSWORD or "DemoPass123!"
DEMO_BOT_TOKEN = settings.DEMO_TELEGRAM_BOT_TOKEN or "TEST_TOKEN_NOT_REAL_REPLACE_BEFORE_USE"

# Shared with generate_landing_demo_asset.py — one source of truth for what the
# extractor is told about this document type.
PROMPT_INSTRUCTIONS = (
    "Extract detailed meeting minutes fields from an internal company meeting "
    "transcript: the client company, location, attendees with their roles, the "
    "agenda actually discussed, decisions reached, and action items — each with "
    "its own owner and due date, in the same order as the action item list."
)

FIELD_SCHEMA = {
    "company_name": {"type": "str", "description": "Client company name", "required": True},
    "location": {"type": "str", "description": "City/office where the meeting took place", "required": True},
    "meeting_date": {"type": "date", "description": "Date the meeting took place, ISO 8601", "required": True},
    "attendees": {"type": "list[str]", "description": "Attendees as 'Full Name — Role'", "required": True},
    "agenda_items": {"type": "list[str]", "description": "Topics discussed, in order", "required": True},
    "summary": {"type": "str", "description": "Narrative summary of the meeting", "required": True},
    "decisions": {"type": "list[str]", "description": "Key decisions agreed", "required": True},
    "action_items": {"type": "list[str]", "description": "Action item descriptions only", "required": True},
    "action_owners": {
        "type": "list[str]",
        "description": "Owner per action item, same order/length as action_items",
        "required": True,
    },
    "action_due_dates": {
        "type": "list[str]",
        "description": "Due date per action item, same order/length as action_items",
        "required": True,
    },
    "next_meeting": {
        "type": "str",
        "description": "Next follow-up date/description, empty string if none mentioned",
        "required": False,
    },
}

_NAVY = RGBColor(0x1F, 0x2A, 0x44)
_GRAY = RGBColor(0x5B, 0x63, 0x70)


def _set_cell_shading(cell: _Cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _heading(document: DocxDocument, text: str, size: int) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = _NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


def _tag_paragraph(document: DocxDocument, tag: str, style: str | None = None) -> None:
    """A paragraph whose only content is a docxtpl control tag ({%p ...%}) — the
    paragraph itself is removed by docxtpl once rendered, only its position matters."""
    document.add_paragraph(tag, style=style)


def _generate_template(output_path: Path) -> None:
    """Builds a valid docxtpl target with zero manual asset creation — a full field_schema
    key becomes a real document section (heading, bullet list, or table), not just a bare
    {{ jinja_tag }} paragraph. Deliberately neutral corporate styling (navy/gray, Calibri),
    not ReportAI's own brand colors — this stands in for *the client's own* template, and
    the product's pitch is "your template, not ours" (see TemplateFidelitySection)."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)
    # Tighter than python-docx's 1" default — this document is dense (title, metadata,
    # 4 sections, a 4-row table) and the default margins pushed the last line to a
    # near-empty second page.
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Title block
    title_p = document.add_paragraph()
    title_run = title_p.add_run("ACTA DE REUNIÓN")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = _NAVY
    subtitle_p = document.add_paragraph()
    subtitle_run = subtitle_p.add_run("{{ company_name }}")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = _GRAY
    subtitle_p.paragraph_format.space_after = Pt(12)

    # Metadata
    meta_p = document.add_paragraph()
    meta_p.add_run("Fecha: ").bold = True
    meta_p.add_run("{{ meeting_date }}    ")
    meta_p.add_run("Lugar: ").bold = True
    meta_p.add_run("{{ location }}")

    # Attendees — real bullet list via docxtpl's paragraph-loop tag
    _heading(document, "Asistentes", 13)
    _tag_paragraph(document, "{%p for a in attendees %}")
    _tag_paragraph(document, "{{ a }}", style="List Bullet")
    _tag_paragraph(document, "{%p endfor %}")

    # Agenda — numbered list, same mechanism
    _heading(document, "Orden del día", 13)
    _tag_paragraph(document, "{%p for item in agenda_items %}")
    _tag_paragraph(document, "{{ item }}", style="List Number")
    _tag_paragraph(document, "{%p endfor %}")

    # Summary
    _heading(document, "Resumen", 13)
    document.add_paragraph("{{ summary }}")

    # Decisions
    _heading(document, "Decisiones adoptadas", 13)
    _tag_paragraph(document, "{%p for d in decisions %}")
    _tag_paragraph(document, "{{ d }}", style="List Bullet")
    _tag_paragraph(document, "{%p endfor %}")

    # Action items — a real 3-column table, header row shaded navy with white text,
    # then a docxtpl row-loop: an opening {%tr for%} marker row, one content row
    # (cloned per action item, indexing the two parallel lists by loop.index0), and
    # a closing {%tr endfor%} marker row — both marker rows are removed by docxtpl.
    _heading(document, "Acciones y responsables", 13)
    table = document.add_table(rows=4, cols=3)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, ["Acción", "Responsable", "Fecha límite"], strict=True):
        _set_cell_shading(cell, "1F2A44")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    table.rows[1].cells[0].paragraphs[0].add_run("{%tr for item in action_items %}")
    content_cells = table.rows[2].cells
    content_cells[0].paragraphs[0].add_run("{{ item }}")
    content_cells[1].paragraphs[0].add_run("{{ action_owners[loop.index0] }}")
    content_cells[2].paragraphs[0].add_run("{{ action_due_dates[loop.index0] }}")
    table.rows[3].cells[0].paragraphs[0].add_run("{%tr endfor %}")

    # Next meeting — optional, only rendered if the extractor found one
    _tag_paragraph(document, "{%p if next_meeting %}")
    next_meeting_p = document.add_paragraph()
    next_meeting_p.add_run("Próxima reunión: ").bold = True
    next_meeting_p.add_run("{{ next_meeting }}")
    _tag_paragraph(document, "{%p endif %}")

    # Footer
    footer_p = document.sections[0].footer.paragraphs[0]
    footer_run = footer_p.add_run(
        "Documento generado automáticamente por ReportAI — confidencial"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = _GRAY

    document.save(str(output_path))


async def _ensure_demo_user(session: AsyncSession, tenant_id: uuid.UUID) -> None:
    user_repo = BaseRepository(TenantUser, session)
    existing = await user_repo.list(filters={"email": DEMO_USER_EMAIL}, limit=1)
    if existing:
        return
    await user_repo.create(
        tenant_id=tenant_id,
        email=DEMO_USER_EMAIL,
        hashed_password=hash_password(DEMO_USER_PASSWORD),
        full_name="Demo Admin",
        role="tenant_admin",
        is_active=True,
    )
    await session.commit()
    print(f"Demo login: {DEMO_USER_EMAIL} / {DEMO_USER_PASSWORD}")


async def seed_demo_tenant(*, reset: bool = False) -> None:
    async with AsyncSessionLocal() as session:
        tenant_repo = BaseRepository(Tenant, session)
        existing = await tenant_repo.list(filters={"slug": "demo"}, limit=1)
        if existing and reset:
            # Every tenant-scoped FK cascades — one DELETE wipes reports, logs,
            # connections, templates and users, then we re-seed from scratch.
            await tenant_repo.delete(existing[0])
            await session.commit()
            print("Demo tenant deleted — re-seeding clean.")
            existing = []
        if existing:
            print("Demo tenant already exists — checking demo user.")
            await _ensure_demo_user(session, existing[0].id)
            return

        tenant = await tenant_repo.create(name="Demo Consulting S.L.", slug="demo", is_active=True)

        doc_type_repo = BaseRepository(DocumentType, session)
        document_type = await doc_type_repo.create(
            tenant_id=tenant.id,
            name="Meeting Minutes",
            description="Internal meeting minutes",
            field_schema=FIELD_SCHEMA,
            prompt_instructions=PROMPT_INSTRUCTIONS,
            # example.com is an RFC 2606 reserved domain (dev only) — a live public demo
            # needs a real monitored inbox so the email leg of the demo actually works.
            notification_emails=[settings.DEMO_NOTIFICATION_EMAIL or "demo-reports@example.com"],
            is_active=True,
        )

        template_path = Path(settings.DOCUMENT_STORAGE_PATH) / "templates" / "demo_meeting_minutes.docx"
        _generate_template(template_path)

        template_repo = BaseRepository(DocumentTemplate, session)
        await template_repo.create(
            tenant_id=tenant.id,
            document_type_id=document_type.id,
            file_path=str(template_path),
            version=1,
            is_active=True,
        )

        connection_repo = BaseRepository(ChannelConnection, session)
        connection = await connection_repo.create(
            tenant_id=tenant.id,
            channel_type="telegram",
            display_name="Demo Telegram Bot",
            credentials={
                "bot_token": DEMO_BOT_TOKEN,
                "secret_token": secrets.token_urlsafe(32),
            },
            is_active=True,
        )

        await session.commit()
        print(f"Seeded demo tenant {tenant.id} with document type {document_type.id}.")
        if not settings.DEMO_TELEGRAM_BOT_TOKEN:
            print("DEMO_TELEGRAM_BOT_TOKEN not set — seeded a fake bot_token (dev only).")
        print(f"Telegram connection {connection.id} — register with scripts/set_telegram_webhook.py")

        await _ensure_demo_user(session, tenant.id)


if __name__ == "__main__":
    asyncio.run(seed_demo_tenant(reset="--reset" in sys.argv[1:]))
